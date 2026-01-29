from __future__ import annotations

from pathlib import Path

from sap_tender_bot.pipeline import attachments as att


def _build_minimal_pdf(text: str) -> bytes:
    text = text.replace("(", "\\(").replace(")", "\\)")
    header = "%PDF-1.4\n"
    objects = []
    objects.append("1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
    objects.append("2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
    objects.append(
        "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n"
    )
    stream_text = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    objects.append(
        f"4 0 obj << /Length {len(stream_text)} >>\n"
        f"stream\n{stream_text}\nendstream\nendobj\n"
    )
    objects.append("5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")

    offsets = [0]
    body = ""
    current = len(header.encode("ascii"))
    for obj in objects:
        offsets.append(current)
        body += obj
        current += len(obj.encode("ascii"))

    xref_start = current
    xref = f"xref\n0 {len(offsets)}\n0000000000 65535 f \n"
    for offset in offsets[1:]:
        xref += f"{offset:010d} 00000 n \n"
    trailer = (
        f"trailer << /Size {len(offsets)} /Root 1 0 R >>\n"
        f"startxref\n{xref_start}\n%%EOF\n"
    )
    return (header + body + xref + trailer).encode("ascii")


def test_pdf_extraction(tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(_build_minimal_pdf("Mandatory Requirements"))
    text, pages = att._extract_pdf_text(pdf_path)
    assert pages == 1
    assert "Mandatory Requirements" in text


def test_docx_extraction(tmp_path: Path) -> None:
    from docx import Document

    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Scope of Work")
    doc.add_paragraph("Deliverables")
    doc.save(doc_path)

    text = att._extract_docx_text(doc_path)
    assert "Scope of Work" in text
    assert "Deliverables" in text


def test_docx_table_extraction(tmp_path: Path) -> None:
    from docx import Document
    from sap_tender_bot.config import AttachmentsConfig

    doc_path = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Requirement"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Response Due"
    table.cell(1, 1).text = "2026-02-15"
    doc.save(doc_path)

    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=False,
        docx_table_max_tables=2,
        docx_table_max_rows=10,
        docx_table_max_cols=5,
        docx_table_char_budget=2000,
    )
    text, _ = att._extract_text(
        doc_path,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        config=config,
    )
    assert "TABLE 1" in text
    assert "Requirement,Value" in text
    assert "Response Due,2026-02-15" in text


def test_docx_table_caps(tmp_path: Path) -> None:
    from docx import Document
    from sap_tender_bot.config import AttachmentsConfig

    doc_path = tmp_path / "table_caps.docx"
    doc = Document()
    table = doc.add_table(rows=4, cols=4)
    for r in range(4):
        for c in range(4):
            table.cell(r, c).text = f"R{r}C{c}"
    doc.save(doc_path)

    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=False,
        docx_table_max_tables=1,
        docx_table_max_rows=2,
        docx_table_max_cols=2,
        docx_table_char_budget=2000,
    )
    text, _ = att._extract_text(
        doc_path,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        config=config,
    )
    lines = [line for line in text.splitlines() if line.strip()]
    assert "TABLE 1" in lines[0]
    assert "R0C0,R0C1" in text
    assert "R0C2" not in text
    assert "R2C0" not in text


def test_spreadsheet_extraction_xlsx(tmp_path: Path) -> None:
    import pandas as pd
    from sap_tender_bot.config import AttachmentsConfig

    xlsx_path = tmp_path / "sample.xlsx"
    df = pd.DataFrame({"Col A": ["A1", "A2"], "Col B": ["B1", "B2"]})
    df.to_excel(xlsx_path, index=False)

    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=False,
        spreadsheets_enabled=True,
        spreadsheet_max_rows=10,
        spreadsheet_max_cols=5,
        spreadsheet_max_sheets=2,
        spreadsheet_sheet_char_budget=2000,
    )
    text, meta = att._extract_text(
        xlsx_path,
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        config=config,
    )
    assert "SHEET:" in text
    assert "Col A" in text
    assert "A1" in text
    assert meta.get("sheet_processed") == 1


def test_spreadsheet_extraction_xls(tmp_path: Path) -> None:
    import pandas as pd
    from sap_tender_bot.config import AttachmentsConfig

    xls_path = tmp_path / "sample.xls"
    df = pd.DataFrame({"Field 1": ["R1", "R2"], "Field 2": ["C1", "C2"]})
    df.to_excel(xls_path, index=False, engine="xlwt")

    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=False,
        spreadsheets_enabled=True,
        spreadsheet_max_rows=10,
        spreadsheet_max_cols=5,
        spreadsheet_max_sheets=2,
        spreadsheet_sheet_char_budget=2000,
    )
    text, meta = att._extract_text(
        xls_path,
        mime_type="application/vnd.ms-excel",
        config=config,
    )
    assert "SHEET:" in text
    assert "Field 1" in text
    assert "R1" in text
    assert meta.get("sheet_processed") == 1


def test_spreadsheet_large_truncation(tmp_path: Path) -> None:
    import pandas as pd
    from sap_tender_bot.config import AttachmentsConfig

    xlsx_path = tmp_path / "large.xlsx"
    df = pd.DataFrame(
        {f"Col {i}": [f"R{r}C{i}" for r in range(30)] for i in range(8)}
    )
    df.to_excel(xlsx_path, index=False)

    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=False,
        spreadsheets_enabled=True,
        spreadsheet_max_rows=5,
        spreadsheet_max_cols=3,
        spreadsheet_max_sheets=1,
        spreadsheet_sheet_char_budget=2000,
    )
    text, meta = att._extract_text(
        xlsx_path,
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        config=config,
    )
    lines = text.splitlines()
    assert lines[0].startswith("SHEET:")
    data_lines = lines[1:]
    assert len(data_lines) <= 6
    assert "Col 0" in text
    assert "Col 3" not in text
    assert meta.get("sheet_processed") == 1


def test_fixture_attachment_extraction_pdf(tmp_path: Path) -> None:
    pdf_text = (
        "Submission Requirements\n"
        "Evaluation Criteria\n"
        "Closing date 2026-03-01"
    )
    pdf_path = tmp_path / "fixture.pdf"
    pdf_path.write_bytes(_build_minimal_pdf(pdf_text))
    text, pages = att._extract_pdf_text(pdf_path)
    assert pages == 1
    assert "Submission Requirements" in text
    assert "Evaluation Criteria" in text
    assert "2026-03-01" in text


def test_fixture_attachment_extraction_docx(tmp_path: Path) -> None:
    from docx import Document
    from sap_tender_bot.config import AttachmentsConfig

    doc_path = tmp_path / "fixture.docx"
    doc = Document()
    doc.add_paragraph("Scope of Work")
    doc.add_paragraph("Evaluation Criteria")
    doc.add_paragraph("Closing date 2026-03-01")
    doc.save(doc_path)

    config = AttachmentsConfig(enabled=True, llm_enabled=False)
    text, _ = att._extract_text(
        doc_path,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        config=config,
    )
    assert "Scope of Work" in text
    assert "Evaluation Criteria" in text
    assert "2026-03-01" in text


def test_fixture_attachment_extraction_xlsx(tmp_path: Path) -> None:
    import pandas as pd
    from sap_tender_bot.config import AttachmentsConfig

    xlsx_path = tmp_path / "fixture.xlsx"
    df = pd.DataFrame({"Milestone": ["Closing date"], "Date": ["2026-03-01"]})
    df.to_excel(xlsx_path, index=False)

    config = AttachmentsConfig(enabled=True, llm_enabled=False, spreadsheets_enabled=True)
    text, _ = att._extract_text(
        xlsx_path,
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        config=config,
    )
    assert "Closing date" in text
    assert "2026-03-01" in text


def test_normalize_summary_evidence_and_dedup() -> None:
    summary = {
        "requirements": [
            "Submit proposal by 2026-02-01 at 2 PM",
            "Submit proposal by 2026-02-01 at 2 PM",
            "Provide a security plan",
        ],
        "dates": ["2026-02-01", "Feb 1, 2026", "TBD"],
        "key_risks": ["Security requirements unclear", "N/A"],
        "attachments_used": ["doc.pdf", "doc.pdf"],
    }
    evidence = "Submission deadline: 2026-02-01 at 2 PM. Provide a security plan."
    normalized = att._normalize_summary(summary, evidence_text=evidence)
    assert normalized["requirements"] == [
        "Submit proposal by 2026-02-01 at 2 PM",
        "Provide a security plan",
    ]
    assert "2026-02-01" in normalized["dates"]
    assert "TBD" not in normalized["dates"]
    assert "Security requirements unclear" in normalized["key_risks"]
    assert "N/A" not in normalized["key_risks"]
    assert normalized["attachments_used"] == ["doc.pdf"]


def test_structured_summary_golden_samples() -> None:
    raw = {
        "submission": [
            {"text": "Submit proposals via portal", "citations": [{"section_heading": "Submission"}]},
            {"text": "Submit proposals via portal", "citations": [{"section_heading": "Submission"}]},
        ],
        "evaluation": [
            {"text": "Evaluation will use weighted scoring", "citations": [{"section_heading": "Evaluation"}]},
        ],
        "scope": [
            {"text": "Implement ERP modernization", "citations": [{"section_heading": "Scope"}]},
        ],
        "deliverables": [
            {"text": "Provide implementation plan", "citations": [{"section_heading": "Deliverables"}]},
        ],
        "schedule": [
            {"text": "Closing date 2026-03-01", "citations": [{"section_heading": "Schedule"}]},
        ],
        "risks": [
            {"text": "Data migration risk", "citations": [{"section_heading": "Risks"}]},
        ],
        "compliance": [
            {"text": "Comply with security policy", "citations": [{"section_heading": "Compliance"}]},
        ],
    }
    evidence = (
        "Submit proposals via portal. Evaluation will use weighted scoring. "
        "Implement ERP modernization. Provide implementation plan. Closing date 2026-03-01. "
        "Data migration risk. Comply with security policy."
    )
    normalized = att._normalize_structured_summary(
        raw, attachment_id="fp123", evidence_text=evidence
    )
    assert len(normalized["submission"]) == 1
    assert len(normalized["evaluation"]) == 1
    assert len(normalized["scope"]) == 1
    assert len(normalized["deliverables"]) == 1
    assert len(normalized["schedule"]) == 1
    assert len(normalized["risks"]) == 1
    assert len(normalized["compliance"]) == 1
    for category in normalized.values():
        for item in category:
            citations = item.get("citations") or []
            assert citations and citations[0].get("attachment_id") == "fp123"


def test_normalize_extracted_text_dedup_lines() -> None:
    repeated = "Standard legal disclaimer applies."
    lines = [f"Line {i}" for i in range(25)]
    text = "\n".join(lines + [repeated] * 3 + ["Tail line"])
    normalized = att._normalize_extracted_text(text)
    assert repeated not in normalized


def test_attachment_caps_per_run_and_tender(tmp_path, monkeypatch) -> None:
    from sap_tender_bot.config import AttachmentsConfig
    from sap_tender_bot.store import TenderStore

    pdf_bytes = _build_minimal_pdf("Mandatory Requirements")

    class FakeResponse:
        def __init__(self, status_code, headers, content):
            self.status_code = status_code
            self.headers = headers
            self._content = content

        def iter_content(self, chunk_size=1024 * 256):
            return iter([self._content])

    def fake_get(url, headers=None, stream=True, timeout=None):
        return FakeResponse(
            200,
            {"Content-Type": "application/pdf", "Content-Length": str(len(pdf_bytes))},
            pdf_bytes,
        )

    monkeypatch.setattr(att.requests, "get", fake_get)

    db_path = tmp_path / "sap_tender.db"
    cache_dir = tmp_path / "cache"
    store = TenderStore(db_path).open()
    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=False,
        max_attachments_per_run=2,
        max_attachments_per_tender=1,
        max_bytes=5_000_000,
    )

    tenders = [
        {"attachments": ["https://example.com/a.pdf", "https://example.com/b.pdf"]},
        {"attachments": ["https://example.com/c.pdf", "https://example.com/d.pdf"]},
    ]
    stats = att.enrich_tenders_with_attachments(
        tenders,
        config=config,
        store=store,
        cache_dir=cache_dir,
    )
    assert stats["attachments_total"] == 2
    store.close()


def test_chunk_building_from_sections() -> None:
    text = (
        "1. Mandatory Requirements\n"
        "Vendor must provide implementation plan.\n"
        "2. Evaluation Criteria\n"
        "Rated criteria include experience.\n"
    )
    _, sections = att._sectionize(text, max_chars=500, max_sections=4)
    chunks = att._build_chunks(text, sections, max_chars=80, overlap_chars=0)
    headings = [c.get("heading") for c in chunks]
    assert "1. Mandatory Requirements" in headings
    assert "2. Evaluation Criteria" in headings
    assert all(c.get("text") for c in chunks)


def test_chunk_selection_keyword_boost(tmp_path: Path) -> None:
    chunks = [
        {"text": "General project overview.", "heading": None},
        {"text": "Evaluation criteria include weighted scoring.", "heading": None},
    ]
    scored = att._score_chunks(
        chunks,
        prompt_embeddings={},
        embeddings_config=None,
        cache_dir=tmp_path,
        attachment_key="unit-test",
    )
    selected = att._select_retrieval_chunks(scored, max_chunks=1)
    assert selected
    assert "Evaluation criteria" in selected[0]["text"]


def test_detect_language_french() -> None:
    text = (
        "Date de clôture: 31 janvier 2026. "
        "Les soumissions doivent inclure les exigences."
    )
    assert att._detect_language(text) == "fr"


def test_detect_language_english() -> None:
    text = "Closing date is January 31, 2026. Submission requirements apply."
    assert att._detect_language(text) == "en"


def test_detect_language_unknown() -> None:
    text = "12345 67890"
    assert att._detect_language(text) is None


def test_scrub_pii_preserves_dates() -> None:
    text = "Email bob@example.com or call 555-123-4567. Deadline 2026-02-01."
    scrubbed = att._scrub_pii(text)
    assert "[REDACTED_EMAIL]" in scrubbed
    assert "[REDACTED_PHONE]" in scrubbed
    assert "2026-02-01" in scrubbed


def test_section_targeting() -> None:
    text = (
        "1. Mandatory Requirements\n"
        "Vendor must provide implementation plan.\n"
        "2. Scope of Work\n"
        "Integrate SAP with CRM.\n"
        "3. Evaluation Criteria\n"
        "Rated criteria include experience.\n"
    )
    excerpt, sections = att._sectionize(text, max_chars=500, max_sections=3)
    headings = [s.get("heading") for s in sections]
    assert "1. Mandatory Requirements" in headings
    assert "2. Scope of Work" in headings
    assert "Evaluation Criteria" in excerpt


def test_section_targeting_french_headings() -> None:
    text = (
        "1. Exigences obligatoires\n"
        "Le fournisseur doit fournir un plan.\n"
        "2. Crit\u00e8res d\u00e9valuation\n"
        "Les soumissions seront notees.\n"
        "3. Date de cl\u00f4ture\n"
        "31 janvier 2026\n"
    )
    excerpt, sections = att._sectionize(text, max_chars=500, max_sections=4)
    headings = [s.get("heading") for s in sections]
    assert "1. Exigences obligatoires" in headings
    assert "2. Crit\u00e8res d\u00e9valuation" in headings
    assert "Date de cl\u00f4ture" in excerpt


def test_section_targeting_fallback_headings() -> None:
    text = (
        "1. Overview\n"
        "General context.\n"
        "2. Pricing and billing\n"
        "Payment terms and schedule.\n"
        "3. Legal terms\n"
        "Standard clauses.\n"
    )
    excerpt, sections = att._sectionize(text, max_chars=500, max_sections=3)
    headings = [s.get("heading") for s in sections]
    assert "1. Overview" in headings
    assert "2. Pricing and billing" in headings
    assert "Legal terms" in excerpt


def test_attachment_cache_reuse(tmp_path, monkeypatch) -> None:
    from sap_tender_bot.config import AttachmentsConfig
    from sap_tender_bot.store import TenderStore

    pdf_bytes = _build_minimal_pdf("Mandatory Requirements")

    class FakeResponse:
        def __init__(self, status_code, headers, content):
            self.status_code = status_code
            self.headers = headers
            self._content = content

        def iter_content(self, chunk_size=1024 * 256):
            if not self._content:
                return iter(())
            return iter([self._content])

    def fake_get(url, headers=None, stream=True, timeout=None):
        if headers and headers.get("If-None-Match"):
            return FakeResponse(304, {}, b"")
        return FakeResponse(
            200,
            {
                "Content-Type": "application/pdf",
                "ETag": '"abc123"',
                "Content-Length": str(len(pdf_bytes)),
            },
            pdf_bytes,
        )

    monkeypatch.setattr(att.requests, "get", fake_get)

    db_path = tmp_path / "sap_tender.db"
    cache_dir = tmp_path / "cache"
    store = TenderStore(db_path).open()
    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=False,
        max_attachments_per_run=5,
        max_attachments_per_tender=2,
        max_bytes=5_000_000,
    )

    tender1 = {"attachments": ["https://example.com/doc.pdf"]}
    stats1 = att.enrich_tenders_with_attachments(
        [tender1],
        config=config,
        store=store,
        cache_dir=cache_dir,
    )
    assert stats1["downloaded"] == 1

    tender2 = {"attachments": ["https://example.com/doc.pdf"]}
    stats2 = att.enrich_tenders_with_attachments(
        [tender2],
        config=config,
        store=store,
        cache_dir=cache_dir,
    )
    assert stats2["downloaded"] == 0
    assert stats2["cached"] >= 1

    store.close()


def test_attachment_cache_ignored_on_amendment(tmp_path, monkeypatch) -> None:
    from sap_tender_bot.config import AttachmentsConfig
    from sap_tender_bot.store import TenderStore

    pdf_bytes = _build_minimal_pdf("Mandatory Requirements")

    class FakeResponse:
        def __init__(self, status_code, headers, content):
            self.status_code = status_code
            self.headers = headers
            self._content = content

        def iter_content(self, chunk_size=1024 * 256):
            if not self._content:
                return iter(())
            return iter([self._content])

    def fake_get(url, headers=None, stream=True, timeout=None):
        if headers and headers.get("If-None-Match"):
            return FakeResponse(304, {}, b"")
        return FakeResponse(
            200,
            {
                "Content-Type": "application/pdf",
                "ETag": '"abc123"',
                "Content-Length": str(len(pdf_bytes)),
                "Last-Modified": "Wed, 01 Jan 2026 00:00:00 GMT",
            },
            pdf_bytes,
        )

    monkeypatch.setattr(att.requests, "get", fake_get)

    call_count = {"count": 0}

    def fake_cohere(prompt, config, logger=None):
        call_count["count"] += 1
        return {
            "requirements": ["Mandatory Requirements"],
            "key_risks": [],
            "dates": [],
            "attachments_used": [],
        }

    monkeypatch.setattr(att, "_cohere_chat", fake_cohere)

    db_path = tmp_path / "sap_tender.db"
    cache_dir = tmp_path / "cache"
    store = TenderStore(db_path).open()
    config = AttachmentsConfig(
        enabled=True,
        llm_enabled=True,
        llm_api_key="test-key",
        max_attachments_per_run=5,
        max_attachments_per_tender=2,
        max_bytes=5_000_000,
    )

    tender1 = {"attachments": ["https://example.com/doc.pdf"]}
    att.enrich_tenders_with_attachments(
        [tender1],
        config=config,
        store=store,
        cache_dir=cache_dir,
    )
    assert call_count["count"] == 1

    tender2 = {"attachments": ["https://example.com/doc.pdf"], "_events": ["amendment"]}
    att.enrich_tenders_with_attachments(
        [tender2],
        config=config,
        store=store,
        cache_dir=cache_dir,
    )
    assert call_count["count"] == 2

    store.close()
