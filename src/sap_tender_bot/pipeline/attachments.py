from __future__ import annotations

import csv
import hashlib
import json
import logging
import random
import re
import shutil
import threading
import time
import unicodedata
import zipfile
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from io import StringIO
from tempfile import TemporaryDirectory
from typing import Any, Iterable, Optional
from urllib.parse import parse_qsl, urlsplit, urlunsplit

import requests

from sap_tender_bot.config import AttachmentsConfig, EmbeddingsConfig
from sap_tender_bot.pipeline.semantic import _cosine_similarity, embed_texts
from sap_tender_bot.store import TenderStore

SUPPORTED_MIME = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/vnd.ms-excel": ".xls",
    "text/plain": ".txt",
    "application/zip": ".zip",
    "application/x-zip-compressed": ".zip",
}
SUPPORTED_EXT = {".pdf", ".docx", ".txt", ".zip", ".xlsx", ".xls"}
ZIP_SUPPORTED_EXT = {".pdf", ".docx", ".txt", ".xlsx", ".xls"}
ZIP_MAX_FILES = 10

BASE_DOWNLOAD_HEADERS = {
    "User-Agent": "sap-tender-bot/attachment/0.1",
    "Accept": "application/pdf,application/octet-stream;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}

BROWSER_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/120.0.0.0 Safari/537.36"
)

BOILERPLATE_REQUIREMENT_TERMS = {
    "title page",
    "table of contents",
    "table of content",
    "toc",
    "abstract",
    "cover page",
    "signature page",
    "bid form",
    "submission form",
    "proposal format",
    "proposal template",
    "response format",
    "pricing",
    "price schedule",
    "price list",
    "schedule of prices",
    "forms",
    "form",
    "glossary",
    "acronyms",
    "definitions",
    "references",
    "report",
    "introduction",
    "work completed",
    "results",
    "conclusion",
    "executive summary",
    "background",
    "overview",
    "purpose",
    "objectives",
    "summary",
    "methodology",
    "approach",
    "scope",
    "timeline",
    "schedule",
    "work plan",
    "project plan",
    "deliverables",
    "milestones",
}

HEADING_TERMS = [
    "mandatory requirements",
    "mandatory criteria",
    "minimum requirements",
    "requirements",
    "technical requirements",
    "functional requirements",
    "compliance requirements",
    "statement of requirements",
    "statement of requirement",
    "scope of work",
    "statement of work",
    "sow",
    "deliverables",
    "mandatory deliverables",
    "minimum deliverables",
    "evaluation criteria",
    "evaluation methodology",
    "evaluation process",
    "evaluation",
    "rated criteria",
    "rated requirements",
    "technical evaluation",
    "submission instructions",
    "submission requirements",
    "proposal requirements",
    "response requirements",
    "proposal submission",
    "submission checklist",
    "bid submission",
    "tender submission",
    "response submission",
    "submission details",
    "vendor response",
    "closing date",
    "closing time",
    "closing date and time",
    "submission deadline",
    "proposal deadline",
    "deadline",
    "key dates",
    "important dates",
    "critical dates",
    "schedule",
    "timeline",
    "project schedule",
    "project timeline",
    "contract term",
    "period of performance",
    "performance period",
    "service levels",
    "service level requirements",
    "security requirements",
    "certification requirements",
    "pricing schedule",
    "pricing requirements",
    "terms and conditions",
    "contractual requirements",
    # French (ASCII-normalized via accent stripping)
    "exigences",
    "exigences obligatoires",
    "exigences minimales",
    "exigences techniques",
    "exigences fonctionnelles",
    "criteres d evaluation",
    "criteres d evaluation technique",
    "criteres d evaluation des offres",
    "criteres obligatoires",
    "criteres minimaux",
    "criteres de selection",
    "instructions de soumission",
    "soumission",
    "soumission des offres",
    "depot des offres",
    "date de cloture",
    "date limite",
    "date limite de soumission",
    "dates importantes",
    "calendrier",
    "echeancier",
    "perimetre des travaux",
    "etendue des travaux",
    "livrables",
    "exigences de securite",
    "exigences de certification",
]

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"\b(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)\d{3}[-.\s]?\d{4}\b")

_RETRIEVAL_PROMPTS = {
    "requirements": (
        "Submission requirements, mandatory requirements, deliverables, scope of work, "
        "technical requirements, compliance requirements."
    ),
    "dates": "Key dates, deadlines, closing date, submission deadline, site visit date.",
    "evaluation": "Evaluation criteria, rated criteria, scoring methodology, selection criteria.",
}

_REQUIREMENT_KEYWORDS = (
    "mandatory",
    "minimum",
    "requirement",
    "submission",
    "deliverable",
    "scope of work",
    "statement of work",
    "compliance",
    "proposal format",
    "response format",
)
_DATE_KEYWORDS = (
    "closing date",
    "deadline",
    "due date",
    "submission date",
    "site visit",
    "key dates",
    "schedule",
    "timeline",
    "important dates",
)
_EVALUATION_KEYWORDS = (
    "evaluation",
    "rated criteria",
    "scoring",
    "weighted",
    "points",
    "selection criteria",
)
_DATE_HINT_RE = re.compile(
    r"\b(?:\d{4}-\d{2}-\d{2}|\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})\b"
)
_CHUNK_MIN_SCORE = 0.06

_FRENCH_STOPWORDS = {
    "le",
    "la",
    "les",
    "des",
    "du",
    "de",
    "et",
    "ou",
    "pour",
    "avec",
    "sans",
    "soumission",
    "exigences",
    "criteres",
    "cloture",
    "date",
    "delai",
}
_FRENCH_ACCENTS_RE = re.compile(r"[àâçéèêëîïôûùüÿœ]")


@dataclass
class AttachmentResult:
    url: str
    url_norm: str
    fingerprint: str
    summary: dict
    requirements: list[dict]
    cached: bool


def _now_utc_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _today_utc() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d")


def _normalize_url(url: str) -> str:
    raw = (url or "").strip()
    if not raw:
        return ""
    parsed = urlsplit(raw)
    scheme = parsed.scheme.lower() or "http"
    host = (parsed.hostname or "").lower()
    port = parsed.port
    netloc = host
    if port:
        netloc = f"{host}:{port}"
    query = "&".join(
        f"{k}={v}" if v != "" else k for k, v in sorted(parse_qsl(parsed.query))
    )
    return urlunsplit((scheme, netloc, parsed.path or "", query, ""))


def _attachment_name(url: str) -> str:
    parsed = urlsplit(url)
    name = Path(parsed.path or "").name
    return name or url


def _default_referer(url: str) -> str:
    parsed = urlsplit(url)
    if parsed.scheme in {"http", "https"} and parsed.hostname:
        return f"{parsed.scheme}://{parsed.hostname}/"
    return ""


def _build_download_headers(url: str, *, user_agent: str, include_referer: bool) -> dict:
    headers = dict(BASE_DOWNLOAD_HEADERS)
    headers["User-Agent"] = user_agent
    if include_referer:
        referer = _default_referer(url)
        if referer:
            headers["Referer"] = referer
    return headers


def _normalize_etag(etag: Optional[str]) -> Optional[str]:
    if not etag:
        return None
    cleaned = str(etag).strip()
    cleaned = cleaned.replace("W/", "").strip()
    return cleaned.strip("\"'")


def _normalize_req_text(text: str) -> str:
    cleaned = re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()
    return re.sub(r"\s+", " ", cleaned)


def _normalize_match_text(text: str) -> str:
    if not text:
        return ""
    cleaned = re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()
    return re.sub(r"\s+", " ", cleaned)


def _dedupe_items(items: list[str]) -> list[str]:
    seen = set()
    deduped = []
    for item in items:
        key = _normalize_match_text(item)
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def _has_evidence_span(item: str, evidence_norm: str) -> bool:
    if not item or not evidence_norm:
        return False
    item_norm = _normalize_match_text(item)
    if not item_norm:
        return False
    if item_norm in evidence_norm:
        return True
    tokens = [t for t in item_norm.split() if len(t) >= 4]
    if not tokens:
        return False
    hits = sum(1 for t in tokens if t in evidence_norm)
    return hits >= min(2, len(tokens))


def _is_low_signal_date(text: str) -> bool:
    if not text:
        return True
    stripped = text.strip()
    if len(stripped) < 4:
        return True
    if not re.search(r"\d", stripped):
        return True
    return False


def _is_low_signal_risk(text: str) -> bool:
    if not text:
        return True
    stripped = text.strip()
    if len(stripped) < 6:
        return True
    words = re.findall(r"[A-Za-z0-9]+", stripped)
    if len(words) <= 2 and len(stripped) < 18:
        return True
    return False


def _is_low_signal_requirement(text: str) -> bool:
    if not text:
        return True
    stripped = text.strip()
    if len(stripped) < 8:
        return True
    words = re.findall(r"[A-Za-z0-9]+", stripped)
    if len(words) <= 2 and len(stripped) < 18:
        return True
    if len(words) <= 4:
        upper_ratio = sum(1 for w in words if w.isupper()) / max(len(words), 1)
        if upper_ratio >= 0.6:
            return True
    return False


def _is_boilerplate_requirement(text: str) -> bool:
    if not text:
        return True
    if _is_low_signal_requirement(text):
        return True
    norm = _normalize_req_text(text)
    if not norm:
        return True
    if norm in BOILERPLATE_REQUIREMENT_TERMS:
        return True
    for prefix in (
        "summary",
        "introduction",
        "conclusion",
        "results",
        "background",
        "overview",
        "purpose",
        "objectives",
        "scope",
        "methodology",
        "approach",
    ):
        if norm.startswith(prefix):
            return True
    if norm.startswith("security requirement"):
        return True
    for prefix in ("appendix", "annex", "annexe"):
        if norm.startswith(prefix):
            return True
    return False


def _guess_extension(url: str, content_type: Optional[str]) -> str:
    url = (url or "").lower()
    for ext in SUPPORTED_EXT:
        if url.endswith(ext):
            return ext
    if content_type:
        mime = content_type.split(";", 1)[0].strip().lower()
        return SUPPORTED_MIME.get(mime, "")
    return ""


def _usage_path(cache_dir: Path, day: str) -> Path:
    return cache_dir / f"usage_{day}.json"


def _load_daily_usage(cache_dir: Path, day: str) -> dict:
    path = _usage_path(cache_dir, day)
    if not path.exists():
        return {"day": day, "count": 0, "updated_at": _now_utc_iso()}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {"day": day, "count": 0, "updated_at": _now_utc_iso()}
    if not isinstance(data, dict):
        return {"day": day, "count": 0, "updated_at": _now_utc_iso()}
    data.setdefault("day", day)
    data.setdefault("count", 0)
    data.setdefault("updated_at", _now_utc_iso())
    return data


def _save_daily_usage(cache_dir: Path, day: str, count: int) -> None:
    cache_dir.mkdir(parents=True, exist_ok=True)
    payload = {"day": day, "count": int(count), "updated_at": _now_utc_iso()}
    _usage_path(cache_dir, day).write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _scrub_pii(text: str) -> str:
    if not text:
        return ""
    text = _EMAIL_RE.sub("[REDACTED_EMAIL]", text)
    text = _PHONE_RE.sub("[REDACTED_PHONE]", text)
    return text


def _scrub_summary_pii(summary: dict) -> dict:
    if not summary:
        return summary
    cleaned = dict(summary)
    cleaned["requirements"] = [_scrub_pii(str(r)) for r in summary.get("requirements", [])]
    cleaned["dates"] = [_scrub_pii(str(d)) for d in summary.get("dates", [])]
    cleaned["key_risks"] = [_scrub_pii(str(r)) for r in summary.get("key_risks", [])]
    cleaned["attachments_used"] = summary.get("attachments_used", [])
    return cleaned


def _scrub_structured_summary(structured: dict) -> dict:
    if not structured or not isinstance(structured, dict):
        return structured
    cleaned: dict = {}
    for key, items in structured.items():
        if not isinstance(items, list):
            cleaned[key] = items
            continue
        scrubbed_items = []
        for item in items:
            if isinstance(item, dict):
                text = _scrub_pii(str(item.get("text") or ""))
                citations = item.get("citations") or []
                scrubbed_items.append({"text": text, "citations": citations})
            else:
                scrubbed_items.append({"text": _scrub_pii(str(item)), "citations": []})
        cleaned[key] = scrubbed_items
    return cleaned


def _extract_json(text: str) -> dict:
    if not text:
        raise ValueError("Empty LLM response")
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in LLM response")
    return json.loads(match.group(0))


def _detect_language(text: str) -> Optional[str]:
    if not text:
        return None
    lowered = text.lower()
    tokens = re.findall(r"[a-zA-Z]+", lowered)
    if not tokens:
        return None
    fr_hits = sum(1 for t in tokens if t in _FRENCH_STOPWORDS)
    accent_hits = bool(_FRENCH_ACCENTS_RE.search(lowered))
    if fr_hits >= 3 or (accent_hits and fr_hits >= 1):
        return "fr"
    return "en"


def _build_prompt(attachment_name: str, excerpt: str, *, language: str = "en") -> str:
    language_note = ""
    if language == "fr":
        language_note = (
            "The attachment is in French. Keep extracted phrases in French.\n"
        )
    return (
        "You are extracting procurement requirements from a tender attachment. "
        "Return ONLY valid JSON with these exact keys:\n"
        "- requirements: array of short strings (exact phrases from the text)\n"
        "- key_risks: array of short strings\n"
        "- dates: array of short strings (closing dates, milestones, site visits)\n"
        "- attachments_used: array of strings (attachment names if present)\n"
        "Rules:\n"
        "- Do NOT invent facts or vendors.\n"
        "- If a requirement is not in the text, omit it.\n"
        "- If nothing is found, return empty arrays but include all keys.\n"
        "No markdown, no extra text.\n\n"
        f"{language_note}"
        f"ATTACHMENT: {attachment_name}\n"
        f"EXCERPT:\n{excerpt}\n"
    )


def _build_repair_prompt(bad_json: str) -> str:
    return (
        "Re-emit the response as VALID JSON ONLY, no extra text. "
        "Use this exact schema and include ALL keys:\n"
        "{\n"
        '  "requirements": [],\n'
        '  "key_risks": [],\n'
        '  "dates": [],\n'
        '  "attachments_used": []\n'
        "}\n"
        "Ensure double quotes and no trailing commas.\n\n"
        f"BAD_JSON:\n{bad_json}\n"
    )


def _build_structured_prompt(
    attachment_name: str,
    attachment_id: str,
    excerpt: str,
    *,
    language: str = "en",
) -> str:
    language_note = ""
    if language == "fr":
        language_note = (
            "The attachment is in French. Keep extracted phrases in French.\n"
        )
    return (
        "You are extracting structured procurement details from a tender attachment. "
        "Return ONLY valid JSON with these exact keys:\n"
        "- submission: array of {text, citations}\n"
        "- evaluation: array of {text, citations}\n"
        "- scope: array of {text, citations}\n"
        "- deliverables: array of {text, citations}\n"
        "- schedule: array of {text, citations}\n"
        "- risks: array of {text, citations}\n"
        "- compliance: array of {text, citations}\n"
        "Each citations entry should be an object with: attachment_id, section_heading, page, offsets.\n"
        "Rules:\n"
        "- Do NOT invent facts.\n"
        "- If a field has no evidence, return an empty array for it.\n"
        "- Use the provided attachment_id.\n"
        "- If section headings are visible, include them in section_heading.\n"
        "No markdown, no extra text.\n\n"
        f"{language_note}"
        f"ATTACHMENT: {attachment_name}\n"
        f"ATTACHMENT_ID: {attachment_id}\n"
        f"EXCERPT:\n{excerpt}\n"
    )


def _cohere_chat_raw(
    prompt: str,
    *,
    api_key: str,
    api_base: str,
    model: str,
    timeout_s: int,
    temperature: float = 0.1,
    max_tokens: int = 360,
    force_v2_only: bool = False,
    logger: Optional[logging.Logger] = None,
) -> str:
    base = api_base.rstrip("/")
    bases = [base]
    if "api.cohere.com" in base:
        bases.append(base.replace("api.cohere.com", "api.cohere.ai"))
    elif "api.cohere.ai" in base:
        bases.append(base.replace("api.cohere.ai", "api.cohere.com"))

    urls: list[str] = []
    for b in bases:
        if b.endswith("/v2"):
            urls.append(f"{b}/chat")
        elif b.endswith("/v1"):
            urls.append(f"{b}/chat")
        else:
            urls.append(f"{b}/v2/chat")
            urls.append(f"{b}/v1/chat")
    urls = list(dict.fromkeys(urls))
    if force_v2_only:
        urls = [u for u in urls if "/v2/" in u]
        if not urls:
            urls = [base.rstrip("/") + "/v2/chat"]

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    schema = {
        "type": "object",
        "required": ["requirements", "key_risks", "dates", "attachments_used"],
        "properties": {
            "requirements": {"type": "array", "items": {"type": "string"}},
            "key_risks": {"type": "array", "items": {"type": "string"}},
            "dates": {"type": "array", "items": {"type": "string"}},
            "attachments_used": {"type": "array", "items": {"type": "string"}},
        },
        "additionalProperties": False,
    }

    payload_v2 = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": max_tokens,
        "response_format": {"type": "json_object", "schema": schema},
    }
    payload_v1 = {
        "model": model,
        "message": prompt,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }

    last_err: Exception | None = None
    for url in urls:
        payload = payload_v2 if "/v2/" in url else payload_v1
        for attempt in range(4):
            try:
                resp = requests.post(url, headers=headers, json=payload, timeout=timeout_s)
                if resp.status_code == 404 and url != urls[-1]:
                    break
                if resp.status_code in {429, 500, 502, 503, 504}:
                    retry_after = resp.headers.get("Retry-After")
                    base_sleep = int(retry_after) if retry_after and retry_after.isdigit() else 1 + attempt
                    sleep_s = base_sleep + random.uniform(0, 0.5)
                    if logger:
                        logger.warning(
                            "Attachment LLM: retryable HTTP %s on %s (attempt %s/4, sleeping %ss)",
                            resp.status_code,
                            url,
                            attempt + 1,
                            round(sleep_s, 2),
                        )
                    time.sleep(sleep_s)
                    continue
                if resp.status_code >= 400:
                    detail = resp.text.strip()
                    msg = f"HTTP {resp.status_code} from {url}"
                    if detail:
                        msg = f"{msg} | {detail[:400]}"
                    raise RuntimeError(msg)
                data = resp.json()
                message = data.get("message")
                if isinstance(message, dict):
                    content = message.get("content")
                    if isinstance(content, list):
                        parts: list[str] = []
                        for item in content:
                            if isinstance(item, dict):
                                if item.get("type") in {"json_object", "json"}:
                                    json_obj = item.get("json")
                                    if isinstance(json_obj, dict):
                                        return json.dumps(json_obj)
                                    if isinstance(json_obj, str) and json_obj.strip():
                                        return json_obj
                                text_item = item.get("text")
                                if isinstance(text_item, str):
                                    parts.append(text_item)
                            elif isinstance(item, str):
                                parts.append(item)
                        if parts:
                            return "".join(parts).strip()
                    elif isinstance(content, str):
                        return content
                text = data.get("text")
                if not text:
                    if isinstance(message, dict):
                        content = message.get("content")
                        if isinstance(content, list):
                            parts = []
                            for item in content:
                                if isinstance(item, dict) and item.get("type") == "text":
                                    parts.append(str(item.get("text", "")))
                                elif isinstance(item, str):
                                    parts.append(item)
                            text = "".join(parts).strip()
                        elif isinstance(content, str):
                            text = content
                if not text and "generations" in data and data["generations"]:
                    text = data["generations"][0].get("text")
                if not text and "message" in data:
                    text = data.get("message")
                return text or ""
            except Exception as exc:
                last_err = exc
                if logger:
                    logger.warning(
                        "Attachment LLM: attempt %s/4 failed for %s (%s)",
                        attempt + 1,
                        url,
                        exc,
                    )
                time.sleep(1 + attempt)

    raise RuntimeError(f"Cohere request failed after retries: {last_err}")


def _cohere_chat(prompt: str, config: AttachmentsConfig, logger=None) -> dict:
    empty = {
        "requirements": [],
        "key_risks": [],
        "dates": [],
        "attachments_used": [],
    }
    try:
        text = _cohere_chat_raw(
            prompt,
            api_key=config.llm_api_key,
            api_base=config.llm_api_base,
            model=config.llm_model,
            timeout_s=config.llm_timeout_s,
            force_v2_only=True,
            logger=logger,
        )
    except Exception as exc:
        if logger:
            logger.warning("Attachment LLM: request failed (%s)", exc)
        empty["_llm_failed"] = True
        return empty
    try:
        return _extract_json(text)
    except Exception as exc:
        if logger:
            logger.warning("Attachment LLM: invalid JSON, attempting repair (%s)", exc)
        repair_prompt = _build_repair_prompt(text)
        try:
            repaired = _cohere_chat_raw(
                repair_prompt,
                api_key=config.llm_api_key,
                api_base=config.llm_api_base,
                model=config.llm_model,
                timeout_s=config.llm_timeout_s,
                force_v2_only=True,
                temperature=0,
                max_tokens=320,
                logger=logger,
            )
            return _extract_json(repaired)
        except Exception as exc2:
            if logger:
                logger.warning("Attachment LLM: repair failed (%s)", exc2)
            strict_prompt = (
                "Return ONLY valid JSON that matches this schema:\n"
                "{\n"
                '  "requirements": [],\n'
                '  "key_risks": [],\n'
                '  "dates": [],\n'
                '  "attachments_used": []\n'
                "}\n"
                "No markdown, no extra text.\n\n"
                f"CONTENT:\n{prompt}\n"
            )
            try:
                strict = _cohere_chat_raw(
                    strict_prompt,
                    api_key=config.llm_api_key,
                    api_base=config.llm_api_base,
                    model=config.llm_model,
                    timeout_s=config.llm_timeout_s,
                    force_v2_only=True,
                    temperature=0,
                    max_tokens=320,
                    logger=logger,
                )
                return _extract_json(strict)
            except Exception as exc3:
                if logger:
                    logger.warning("Attachment LLM: strict JSON failed (%s)", exc3)
                empty["_llm_failed"] = True
                return empty


def _cohere_chat_structured(prompt: str, config: AttachmentsConfig, logger=None) -> dict:
    empty = {
        "submission": [],
        "evaluation": [],
        "scope": [],
        "deliverables": [],
        "schedule": [],
        "risks": [],
        "compliance": [],
    }
    base = config.llm_api_base.rstrip("/")
    bases = [base]
    if "api.cohere.com" in base:
        bases.append(base.replace("api.cohere.com", "api.cohere.ai"))
    elif "api.cohere.ai" in base:
        bases.append(base.replace("api.cohere.ai", "api.cohere.com"))

    urls: list[str] = []
    for b in bases:
        if b.endswith("/v2"):
            urls.append(f"{b}/chat")
        elif b.endswith("/v1"):
            urls.append(f"{b}/chat")
        else:
            urls.append(f"{b}/v2/chat")
            urls.append(f"{b}/v1/chat")
    urls = list(dict.fromkeys(urls))

    headers = {
        "Authorization": f"Bearer {config.llm_api_key}",
        "Content-Type": "application/json",
    }
    item_schema = {
        "type": "object",
        "required": ["text", "citations"],
        "properties": {
            "text": {"type": "string"},
            "citations": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "attachment_id": {"type": "string"},
                        "section_heading": {"type": "string"},
                        "page": {"type": ["integer", "null"]},
                        "offsets": {
                            "type": ["array", "null"],
                            "items": {"type": "integer"},
                        },
                    },
                    "additionalProperties": False,
                },
            },
        },
        "additionalProperties": False,
    }
    schema = {
        "type": "object",
        "required": [
            "submission",
            "evaluation",
            "scope",
            "deliverables",
            "schedule",
            "risks",
            "compliance",
        ],
        "properties": {
            "submission": {"type": "array", "items": item_schema},
            "evaluation": {"type": "array", "items": item_schema},
            "scope": {"type": "array", "items": item_schema},
            "deliverables": {"type": "array", "items": item_schema},
            "schedule": {"type": "array", "items": item_schema},
            "risks": {"type": "array", "items": item_schema},
            "compliance": {"type": "array", "items": item_schema},
        },
        "additionalProperties": False,
    }
    payload_v2 = {
        "model": config.llm_model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 520,
        "response_format": {"type": "json_object", "schema": schema},
    }
    payload_v1 = {
        "model": config.llm_model,
        "message": prompt,
        "temperature": 0.1,
        "max_tokens": 520,
    }

    last_err: Exception | None = None
    for url in urls:
        payload = payload_v2 if "/v2/" in url else payload_v1
        for attempt in range(4):
            try:
                resp = requests.post(url, headers=headers, json=payload, timeout=config.llm_timeout_s)
                if resp.status_code == 404 and url != urls[-1]:
                    break
                if resp.status_code in {429, 500, 502, 503, 504}:
                    retry_after = resp.headers.get("Retry-After")
                    sleep_s = int(retry_after) if retry_after and retry_after.isdigit() else 1 + attempt
                    time.sleep(sleep_s)
                    continue
                if resp.status_code >= 400:
                    detail = resp.text.strip()
                    msg = f"HTTP {resp.status_code} from {url}"
                    if detail:
                        msg = f"{msg} | {detail[:400]}"
                    raise RuntimeError(msg)
                data = resp.json()
                message = data.get("message")
                if isinstance(message, dict):
                    content = message.get("content")
                    if isinstance(content, list):
                        parts: list[str] = []
                        for item in content:
                            if isinstance(item, dict):
                                if item.get("type") in {"json_object", "json"}:
                                    json_obj = item.get("json")
                                    if isinstance(json_obj, dict):
                                        return json_obj
                                    if isinstance(json_obj, str) and json_obj.strip():
                                        return json.loads(json_obj)
                                text_item = item.get("text")
                                if isinstance(text_item, str):
                                    parts.append(text_item)
                            elif isinstance(item, str):
                                parts.append(item)
                        if parts:
                            return json.loads("".join(parts).strip())
                    elif isinstance(content, str):
                        return json.loads(content)
                text = data.get("text")
                if not text and "generations" in data and data["generations"]:
                    text = data["generations"][0].get("text")
                if not text and "message" in data:
                    text = data.get("message")
                if text:
                    return json.loads(text)
                return empty
            except Exception as exc:
                last_err = exc
                time.sleep(1 + attempt)
    if logger:
        logger.warning("Attachment structured LLM failed: %s", last_err)
    return empty


def _coerce_list(value: Any, max_items: int = 8) -> list[str]:
    if value is None:
        items: list[str] = []
    elif isinstance(value, list):
        items = [str(v).strip() for v in value if str(v).strip()]
    elif isinstance(value, str):
        items = [value.strip()] if value.strip() else []
    else:
        items = [str(value).strip()] if str(value).strip() else []
    return items[:max_items]


def _normalize_summary(summary: dict, *, evidence_text: str = "") -> dict:
    evidence_norm = _normalize_match_text(evidence_text)
    requirements = []
    for r in _coerce_list(summary.get("requirements"), max_items=8):
        if _is_boilerplate_requirement(r):
            continue
        if evidence_norm and not _has_evidence_span(r, evidence_norm):
            continue
        requirements.append(r)
    requirements = _dedupe_items(requirements)

    dates = []
    for d in _coerce_list(summary.get("dates"), max_items=6):
        if _is_low_signal_date(d):
            continue
        if evidence_norm and not _has_evidence_span(d, evidence_norm):
            continue
        dates.append(d)
    dates = _dedupe_items(dates)

    risks = []
    for r in _coerce_list(summary.get("key_risks"), max_items=6):
        if _is_low_signal_risk(r):
            continue
        risks.append(r)
    risks = _dedupe_items(risks)

    attachments_used = _dedupe_items(
        _coerce_list(summary.get("attachments_used"), max_items=4)
    )

    normalized = {
        "requirements": requirements,
        "key_risks": risks,
        "dates": dates,
        "attachments_used": attachments_used,
    }
    return normalized


def _normalize_citations(
    citations: Any,
    *,
    attachment_id: str,
    default_heading: str = "",
) -> list[dict]:
    normalized: list[dict] = []
    if isinstance(citations, list):
        source = citations
    elif citations is None:
        source = []
    else:
        source = [citations]
    for item in source:
        if isinstance(item, dict):
            attachment = str(item.get("attachment_id") or attachment_id or "")
            heading = str(item.get("section_heading") or default_heading or "")
            page = item.get("page")
            if page is not None:
                try:
                    page = int(page)
                except Exception:
                    page = None
            offsets = item.get("offsets")
            if isinstance(offsets, list) and len(offsets) >= 2:
                try:
                    offsets = [int(offsets[0]), int(offsets[1])]
                except Exception:
                    offsets = None
            else:
                offsets = None
            normalized.append(
                {
                    "attachment_id": attachment,
                    "section_heading": heading,
                    "page": page,
                    "offsets": offsets,
                }
            )
        elif isinstance(item, str):
            normalized.append(
                {
                    "attachment_id": attachment_id,
                    "section_heading": item,
                    "page": None,
                    "offsets": None,
                }
            )
    if not normalized and attachment_id:
        normalized.append(
            {
                "attachment_id": attachment_id,
                "section_heading": default_heading,
                "page": None,
                "offsets": None,
            }
        )
    return normalized


def _normalize_structured_items(
    items: Any,
    *,
    attachment_id: str,
    evidence_text: str,
    low_signal_check,
) -> list[dict]:
    normalized: list[dict] = []
    if items is None:
        raw_items: list[Any] = []
    elif isinstance(items, list):
        raw_items = items
    else:
        raw_items = [items]
    for item in raw_items:
        if isinstance(item, dict):
            text = str(item.get("text") or "").strip()
            citations = _normalize_citations(
                item.get("citations"),
                attachment_id=attachment_id,
                default_heading=str(item.get("section_heading") or ""),
            )
        else:
            text = str(item).strip()
            citations = _normalize_citations(
                [],
                attachment_id=attachment_id,
                default_heading="",
            )
        if not text or low_signal_check(text):
            continue
        if evidence_text and not _has_evidence_span(text, _normalize_match_text(evidence_text)):
            continue
        normalized.append({"text": text, "citations": citations})
    normalized = _dedupe_structured(normalized)
    return normalized


def _dedupe_structured(items: list[dict]) -> list[dict]:
    seen = set()
    deduped = []
    for item in items:
        text = str(item.get("text") or "").strip()
        key = _normalize_match_text(text)
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def _normalize_structured_summary(
    summary: dict,
    *,
    attachment_id: str,
    evidence_text: str,
) -> dict:
    if not isinstance(summary, dict):
        summary = {}
    return {
        "submission": _normalize_structured_items(
            summary.get("submission"),
            attachment_id=attachment_id,
            evidence_text=evidence_text,
            low_signal_check=_is_low_signal_requirement,
        ),
        "evaluation": _normalize_structured_items(
            summary.get("evaluation"),
            attachment_id=attachment_id,
            evidence_text=evidence_text,
            low_signal_check=_is_low_signal_requirement,
        ),
        "scope": _normalize_structured_items(
            summary.get("scope"),
            attachment_id=attachment_id,
            evidence_text=evidence_text,
            low_signal_check=_is_low_signal_requirement,
        ),
        "deliverables": _normalize_structured_items(
            summary.get("deliverables"),
            attachment_id=attachment_id,
            evidence_text=evidence_text,
            low_signal_check=_is_low_signal_requirement,
        ),
        "schedule": _normalize_structured_items(
            summary.get("schedule"),
            attachment_id=attachment_id,
            evidence_text=evidence_text,
            low_signal_check=_is_low_signal_date,
        ),
        "risks": _normalize_structured_items(
            summary.get("risks"),
            attachment_id=attachment_id,
            evidence_text=evidence_text,
            low_signal_check=_is_low_signal_risk,
        ),
        "compliance": _normalize_structured_items(
            summary.get("compliance"),
            attachment_id=attachment_id,
            evidence_text=evidence_text,
            low_signal_check=_is_low_signal_requirement,
        ),
    }


def _table_to_csv_text(table, max_rows: int, max_cols: int) -> str:
    buffer = StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    for row_idx, row in enumerate(table.rows):
        if max_rows > 0 and row_idx >= max_rows:
            break
        cells = row.cells
        if max_cols > 0:
            cells = cells[:max_cols]
        values = []
        for cell in cells:
            text = " ".join(str(cell.text or "").split())
            values.append(text)
        writer.writerow(values)
    return buffer.getvalue().strip()


def _extract_docx_text(
    path: Path,
    *,
    config: Optional[AttachmentsConfig] = None,
    logger: Optional[logging.Logger] = None,
) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for DOCX extraction") from exc
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    if not (config is None or getattr(config, "docx_tables_enabled", True)):
        return "\n".join(parts)

    max_tables = int(getattr(config, "docx_table_max_tables", 0) or 0)
    max_rows = int(getattr(config, "docx_table_max_rows", 0) or 0)
    max_cols = int(getattr(config, "docx_table_max_cols", 0) or 0)
    table_char_budget = int(getattr(config, "docx_table_char_budget", 0) or 0)

    try:
        for idx, table in enumerate(doc.tables):
            if max_tables > 0 and idx >= max_tables:
                break
            table_text = _table_to_csv_text(table, max_rows, max_cols)
            if table_char_budget > 0 and len(table_text) > table_char_budget:
                table_text = table_text[:table_char_budget].rsplit("\n", 1)[0].strip()
            if table_text:
                parts.append(f"TABLE {idx + 1}\n{table_text}")
    except Exception as exc:
        if logger:
            logger.warning("DOCX table extraction failed for %s: %s", path.name, exc)
    return "\n".join(parts)


def _extract_pdf_text(path: Path) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("pypdf is required for PDF extraction") from exc
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    cleaned_pages = _strip_repeated_lines(pages)
    return "\n\n".join(cleaned_pages).strip(), len(reader.pages)


def _extract_txt_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _dataframe_to_csv_text(df) -> str:
    if df is None:
        return ""
    try:
        columns = [str(c) for c in df.columns.tolist()]
    except Exception:
        columns = []
    buffer = StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    if columns:
        writer.writerow(columns)
    for row in getattr(df, "itertuples")(index=False, name=None):
        cleaned = ["" if v is None else str(v) for v in row]
        writer.writerow(cleaned)
    return buffer.getvalue().strip()


def _extract_spreadsheet_text(
    path: Path,
    *,
    config: AttachmentsConfig,
    logger: Optional[logging.Logger] = None,
) -> tuple[str, dict]:
    if not config.spreadsheets_enabled:
        return "", {"spreadsheet_skipped": True, "sheet_count": 0}
    max_bytes = int(getattr(config, "spreadsheet_max_bytes", 0) or 0)
    if max_bytes > 0 and path.exists() and path.stat().st_size > max_bytes:
        return "", {"spreadsheet_too_large": True, "sheet_count": 0}
    try:
        import pandas as pd
    except Exception as exc:
        raise RuntimeError("pandas is required for spreadsheet extraction") from exc

    engine = None
    if path.suffix.lower() == ".xlsx":
        engine = "openpyxl"
    elif path.suffix.lower() == ".xls":
        engine = "xlrd"

    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str, engine=engine)
    except Exception as exc:
        if logger:
            logger.warning("Spreadsheet extraction failed for %s: %s", path.name, exc)
        return "", {"spreadsheet_failed": True, "sheet_count": 0}

    max_rows = int(getattr(config, "spreadsheet_max_rows", 0) or 0)
    max_cols = int(getattr(config, "spreadsheet_max_cols", 0) or 0)
    max_sheets = int(getattr(config, "spreadsheet_max_sheets", 0) or 0)
    sheet_char_budget = int(getattr(config, "spreadsheet_sheet_char_budget", 0) or 0)

    sheet_names = list(sheets.keys())
    if max_sheets > 0:
        sheet_names = sheet_names[:max_sheets]

    text_parts: list[str] = []
    processed = 0
    for name in sheet_names:
        df = sheets.get(name)
        if df is None:
            continue
        if max_rows > 0:
            df = df.head(max_rows)
        if max_cols > 0 and getattr(df, "shape", (0, 0))[1] > max_cols:
            df = df.iloc[:, :max_cols]
        df = df.fillna("")
        sheet_text = _dataframe_to_csv_text(df)
        if sheet_char_budget > 0 and len(sheet_text) > sheet_char_budget:
            sheet_text = sheet_text[:sheet_char_budget].rsplit("\n", 1)[0].strip()
        if sheet_text:
            text_parts.append(f"SHEET: {name}\n{sheet_text}".strip())
            processed += 1

    meta = {
        "sheet_count": len(sheets),
        "sheet_processed": processed,
        "spreadsheet_rows": max_rows,
        "spreadsheet_cols": max_cols,
    }
    return "\n\n".join(text_parts).strip(), meta


def _extract_zip_text(
    path: Path,
    *,
    config: AttachmentsConfig,
    logger: Optional[logging.Logger] = None,
) -> tuple[str, dict]:
    total_bytes = 0
    extracted = 0
    page_count = 0
    text_parts: list[str] = []
    try:
        with zipfile.ZipFile(path) as zf, TemporaryDirectory(dir=path.parent) as tmp_dir:
            for info in zf.infolist():
                if extracted >= ZIP_MAX_FILES:
                    break
                if info.is_dir():
                    continue
                ext = Path(info.filename).suffix.lower()
                if ext not in ZIP_SUPPORTED_EXT:
                    continue
                size = int(info.file_size or 0)
                if size <= 0:
                    continue
                if size > config.max_bytes:
                    continue
                if total_bytes + size > config.max_bytes:
                    continue

                safe_name = f"{extracted:02d}_{Path(info.filename).name}"
                tmp_path = Path(tmp_dir) / safe_name
                with zf.open(info) as src, tmp_path.open("wb") as dst:
                    shutil.copyfileobj(src, dst)
                total_bytes += size
                extracted += 1

                if ext == ".pdf":
                    text, pages = _extract_pdf_text(tmp_path)
                    page_count += pages
                elif ext == ".docx":
                    text = _extract_docx_text(tmp_path, config=config, logger=logger)
                elif ext in {".xlsx", ".xls"}:
                    text, sheet_meta = _extract_spreadsheet_text(
                        tmp_path, config=config, logger=logger
                    )
                    page_count += int(sheet_meta.get("sheet_processed") or 0)
                else:
                    text = _extract_txt_text(tmp_path)

                if text:
                    text_parts.append(f"FILE: {info.filename}\n{text}")
    except Exception as exc:
        if logger:
            logger.warning("ZIP extraction failed for %s: %s", path.name, exc)
        return "", {"page_count": 0, "zip_files": 0, "zip_bytes": 0}

    meta = {"page_count": page_count, "zip_files": extracted, "zip_bytes": total_bytes}
    return "\n\n".join(text_parts).strip(), meta


def _ocr_pdf_text(path: Path, logger: Optional[logging.Logger] = None) -> str:
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except Exception:
        if logger:
            logger.warning("OCR dependencies missing; install pdf2image + pytesseract to enable OCR.")
        return ""
    try:
        images = convert_from_path(str(path), dpi=200)
        text_parts = [pytesseract.image_to_string(img) for img in images]
        return "\n".join(part for part in text_parts if part).strip()
    except Exception as exc:
        if logger:
            logger.warning("OCR failed for %s: %s", path.name, exc)
        return ""


def _strip_repeated_lines(pages: list[str]) -> list[str]:
    if len(pages) < 2:
        return [p.strip() for p in pages]
    counts: Counter[str] = Counter()
    for page in pages:
        for line in page.splitlines():
            line = line.strip()
            if line:
                counts[line] += 1
    repeated = {line for line, count in counts.items() if count >= 2}
    cleaned_pages = []
    for page in pages:
        lines = [line for line in page.splitlines() if line.strip() and line.strip() not in repeated]
        cleaned_pages.append("\n".join(lines).strip())
    return cleaned_pages


def _normalize_extracted_text(text: str) -> str:
    if not text:
        return ""
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        lines.append(line)
    if len(lines) >= 30:
        counts: Counter[str] = Counter(lines)
        repeated = {
            line
            for line, count in counts.items()
            if count >= 3 and 12 <= len(line) <= 180
        }
        if repeated:
            lines = [line for line in lines if line not in repeated]
    cleaned = "\n".join(lines)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def _strip_accents(text: str) -> str:
    if not text:
        return ""
    return "".join(
        ch for ch in unicodedata.normalize("NFKD", text) if not unicodedata.combining(ch)
    )


def _looks_like_heading(line: str) -> bool:
    if not line:
        return False
    if len(line) > 120:
        return False
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", line):
        return True
    if line.endswith(":"):
        return True
    if line.isupper() and len(line) > 4:
        return True
    return False


def _chunk_text_by_chars(text: str, max_chars: int, overlap_chars: int) -> list[dict]:
    if not text:
        return []
    if max_chars <= 0:
        return [{"text": text.strip(), "start": 0, "end": len(text)}]
    chunks: list[dict] = []
    start = 0
    length = len(text)
    overlap = max(0, overlap_chars)
    while start < length:
        end = min(length, start + max_chars)
        slice_text = text[start:end]
        split_at = max(slice_text.rfind("\n\n"), slice_text.rfind("\n"), slice_text.rfind(" "))
        if split_at > int(max_chars * 0.6):
            end = start + split_at
            slice_text = text[start:end]
        cleaned = slice_text.strip()
        if cleaned:
            chunks.append({"text": cleaned, "start": start, "end": end})
        if end >= length:
            break
        start = max(0, end - overlap)
        if start >= length:
            break
    return chunks


def _build_chunks(
    text: str,
    sections: list[dict],
    *,
    max_chars: int,
    overlap_chars: int,
) -> list[dict]:
    chunks: list[dict] = []
    if sections:
        for section_index, section in enumerate(sections):
            body = str(section.get("text") or "").strip()
            if not body:
                continue
            for idx, chunk in enumerate(_chunk_text_by_chars(body, max_chars, overlap_chars)):
                chunks.append(
                    {
                        "text": chunk["text"],
                        "heading": section.get("heading"),
                        "section_index": section_index,
                        "chunk_index": idx,
                        "start_offset": chunk["start"],
                        "end_offset": chunk["end"],
                        "source": "section",
                    }
                )
    else:
        for idx, chunk in enumerate(_chunk_text_by_chars(text, max_chars, overlap_chars)):
            chunks.append(
                {
                    "text": chunk["text"],
                    "heading": None,
                    "section_index": None,
                    "chunk_index": idx,
                    "start_offset": chunk["start"],
                    "end_offset": chunk["end"],
                    "source": "text",
                }
            )
    return chunks


def _keyword_boost(text: str) -> float:
    lowered = text.lower()
    boost = 0.0
    if any(term in lowered for term in _REQUIREMENT_KEYWORDS):
        boost += 0.08
    if any(term in lowered for term in _DATE_KEYWORDS) or _DATE_HINT_RE.search(text):
        boost += 0.07
    if any(term in lowered for term in _EVALUATION_KEYWORDS):
        boost += 0.06
    return boost


def _score_chunks(
    chunks: list[dict],
    *,
    prompt_embeddings: dict[str, list[float]] | None,
    embeddings_config: EmbeddingsConfig | None,
    cache_dir: Path,
    attachment_key: str,
    logger: Optional[logging.Logger] = None,
) -> list[dict]:
    if not chunks:
        return []
    prompt_embeddings = prompt_embeddings or {}
    use_embeddings = bool(prompt_embeddings) and embeddings_config is not None

    chunk_embeddings: dict[str, list[float]] = {}
    if use_embeddings:
        texts_by_key = {}
        for idx, chunk in enumerate(chunks):
            key = f"attachment:{attachment_key}:chunk:{idx}"
            chunk["chunk_id"] = key
            texts_by_key[key] = str(chunk.get("text") or "")
        chunk_embeddings, embedded = embed_texts(
            texts_by_key, embeddings_config, cache_dir, logger=logger
        )
        use_embeddings = embedded and bool(chunk_embeddings)

    for idx, chunk in enumerate(chunks):
        base_score = 0.0
        best_prompt = ""
        if use_embeddings:
            key = chunk.get("chunk_id") or f"attachment:{attachment_key}:chunk:{idx}"
            chunk_vec = chunk_embeddings.get(key)
            if chunk_vec:
                for prompt_name, prompt_vec in prompt_embeddings.items():
                    sim = _cosine_similarity(chunk_vec, prompt_vec)
                    if sim > base_score:
                        base_score = sim
                        best_prompt = prompt_name
        boost = _keyword_boost(str(chunk.get("text") or ""))
        total = min(1.0, base_score + boost)
        chunk["score"] = round(total, 4)
        chunk["score_prompt"] = best_prompt or None
        chunk["score_boost"] = round(boost, 4)
        chunk["score_similarity"] = round(base_score, 4)
        chunk["used_embeddings"] = bool(use_embeddings)
    return chunks


def _select_retrieval_chunks(chunks: list[dict], max_chunks: int) -> list[dict]:
    if not chunks or max_chunks <= 0:
        return []
    ranked = [c for c in chunks if (c.get("score") or 0) >= _CHUNK_MIN_SCORE]
    if not ranked:
        ranked = chunks
    ranked.sort(key=lambda c: (c.get("score") or 0), reverse=True)
    return ranked[:max_chunks]


def _build_retrieval_excerpt(chunks: list[dict], max_chars: int) -> str:
    parts: list[str] = []
    total = 0
    for idx, chunk in enumerate(chunks, start=1):
        heading = _trim_heading(chunk.get("heading"), max_len=80)
        label = f"CHUNK {idx}"
        if heading:
            label = f"{label} ({heading})"
        chunk_text = str(chunk.get("text") or "").strip()
        if not chunk_text:
            continue
        block = f"{label}\n{chunk_text}".strip()
        if max_chars > 0 and total + len(block) > max_chars:
            remaining = max_chars - total
            if remaining <= 20:
                break
            block = block[:remaining].rsplit("\n", 1)[0].strip()
        if block:
            parts.append(block)
            total += len(block) + 2
        if max_chars > 0 and total >= max_chars:
            break
    return "\n\n".join(parts).strip()


def _sectionize(text: str, max_chars: int, max_sections: int) -> tuple[str, list[dict]]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    headings: list[int] = []
    for idx, line in enumerate(lines):
        lowered = line.lower()
        lowered_norm = _strip_accents(lowered)
        if not any(term in lowered_norm for term in HEADING_TERMS):
            continue
        if _looks_like_heading(line):
            headings.append(idx)

    sections: list[dict] = []
    if not headings:
        fallback = []
        for idx, line in enumerate(lines):
            if not _looks_like_heading(line):
                continue
            if line.endswith("."):
                continue
            word_count = len(line.split())
            if word_count <= 0 or word_count > 12:
                continue
            fallback.append(idx)
        headings = fallback[:max_sections]
        if not headings:
            excerpt = text[:max_chars].strip()
            return excerpt, []

    for i, h_idx in enumerate(headings[:max_sections]):
        start = h_idx
        end = headings[i + 1] if i + 1 < len(headings) else len(lines)
        heading = lines[h_idx]
        body = "\n".join(lines[start + 1 : end]).strip()
        sections.append(
            {
                "heading": heading,
                "text": body,
                "chars": len(body),
                "start_line": start,
                "end_line": end,
            }
        )

    assembled_parts = []
    for section in sections:
        if section["text"]:
            assembled_parts.append(f"{section['heading']}\n{section['text']}")
        else:
            assembled_parts.append(section["heading"])
    assembled = "\n\n".join(assembled_parts)
    if len(assembled) > max_chars:
        assembled = assembled[:max_chars].rsplit("\n", 1)[0].strip()
    return assembled, sections


def _trim_heading(text: str, max_len: int = 80) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "").strip())
    if not cleaned:
        return ""
    if len(cleaned) <= max_len:
        return cleaned
    truncated = cleaned[:max_len].rsplit(" ", 1)[0].strip()
    return truncated or cleaned[:max_len].strip()


def _truncate_sentence(text: str, max_len: int) -> str:
    if len(text) <= max_len:
        return text
    truncated = text[:max_len].rsplit(" ", 1)[0].strip()
    return truncated or text[:max_len].strip()


def _build_attachment_memo(
    summary: dict,
    *,
    attachment_name: str = "",
    source_sections: Optional[list[str]] = None,
) -> str:
    if not summary:
        return ""
    reqs = summary.get("requirements") or []
    dates = summary.get("dates") or []
    risks = summary.get("key_risks") or []
    parts: list[str] = []
    if attachment_name:
        parts.append(f"Attachment: {attachment_name}.")
    if source_sections:
        cleaned = []
        for section in source_sections:
            trimmed = _trim_heading(section)
            if trimmed:
                cleaned.append(trimmed)
            if len(cleaned) >= 2:
                break
        if cleaned:
            parts.append(f"Sections: {', '.join(cleaned)}.")
    if reqs:
        parts.append(f"Requirements: {', '.join(str(r) for r in reqs[:2])}.")
    if dates:
        parts.append(f"Dates: {', '.join(str(d) for d in dates[:2])}.")
    if risks:
        parts.append(f"Risks: {', '.join(str(r) for r in risks[:2])}.")
    memo = " ".join(p.strip() for p in parts if p).strip()
    return _truncate_sentence(memo, 320)


def _excerpt_fallback_memo(excerpt: str, *, attachment_name: str = "") -> str:
    if not excerpt:
        return ""
    prefix = f"Attachment: {attachment_name}. " if attachment_name else ""
    cleaned = " ".join(excerpt.split())
    budget = max(0, 300 - len(prefix) - len("Doc excerpt: "))
    if budget and len(cleaned) > budget:
        cleaned = cleaned[:budget].rsplit(" ", 1)[0].strip()
    if not cleaned:
        return ""
    return f"{prefix}Doc excerpt: {cleaned}"


def _download_attachment(
    url: str,
    *,
    config: AttachmentsConfig,
    store: TenderStore,
    cache_dir: Path,
    logger: Optional[logging.Logger] = None,
) -> tuple[Optional[dict], Optional[Path], Optional[str], bool]:
    url_norm = _normalize_url(url)
    if not url_norm:
        return None, None, "invalid_url", False

    cache_dir.mkdir(parents=True, exist_ok=True)
    latest = store.fetch_latest_attachment_meta(url_norm=url_norm)
    headers = _build_download_headers(url, user_agent=BASE_DOWNLOAD_HEADERS["User-Agent"], include_referer=False)
    if latest:
        if latest.get("etag"):
            headers["If-None-Match"] = latest.get("etag")
        if latest.get("last_modified"):
            headers["If-Modified-Since"] = latest.get("last_modified")

    def _retry_delay(attempt: int) -> None:
        time.sleep(1 + attempt)

    used_browser_fallback = False
    for attempt in range(config.download_retries + 1):
        try:
            resp = requests.get(url, headers=headers, stream=True, timeout=config.timeout_s)
            if resp.status_code == 304 and latest:
                cached = store.fetch_attachment_cache_by_meta(
                    url_norm=url_norm,
                    etag=latest.get("etag"),
                    last_modified=latest.get("last_modified"),
                )
                if cached and cached.get("text_path") and Path(cached["text_path"]).exists():
                    return cached, Path(cached["text_path"]), None, True
                headers.pop("If-None-Match", None)
                headers.pop("If-Modified-Since", None)
                resp = requests.get(url, headers=headers, stream=True, timeout=config.timeout_s)

            if resp.status_code in {401, 403} and not used_browser_fallback:
                used_browser_fallback = True
                headers = _build_download_headers(url, user_agent=BROWSER_UA, include_referer=True)
                if latest:
                    if latest.get("etag"):
                        headers["If-None-Match"] = latest.get("etag")
                    if latest.get("last_modified"):
                        headers["If-Modified-Since"] = latest.get("last_modified")
                if logger:
                    logger.warning(
                        "Attachment download HTTP %s for %s; retrying with browser headers",
                        resp.status_code,
                        url,
                    )
                resp = requests.get(url, headers=headers, stream=True, timeout=config.timeout_s)

            if resp.status_code >= 400:
                raise RuntimeError(f"HTTP {resp.status_code}")

            content_type = resp.headers.get("Content-Type", "")
            extension = _guess_extension(url, content_type)
            if extension not in SUPPORTED_EXT:
                return None, None, "unsupported_type", False

            content_length = resp.headers.get("Content-Length")
            if content_length and content_length.isdigit():
                if int(content_length) > config.max_bytes:
                    return None, None, "too_large", False

            tmp_path = cache_dir / f"tmp_{hashlib.sha256(url_norm.encode('utf-8')).hexdigest()}.bin"
            digest = hashlib.sha256()
            size = 0
            with tmp_path.open("wb") as fp:
                for chunk in resp.iter_content(chunk_size=1024 * 256):
                    if not chunk:
                        continue
                    size += len(chunk)
                    if size > config.max_bytes:
                        fp.close()
                        tmp_path.unlink(missing_ok=True)
                        return None, None, "too_large", False
                    digest.update(chunk)
                    fp.write(chunk)

            content_hash = digest.hexdigest()
            etag = _normalize_etag(resp.headers.get("ETag"))
            last_modified = resp.headers.get("Last-Modified")
            filename = _attachment_name(url)
            if etag or last_modified:
                fingerprint_seed = (
                    f"{url_norm}|{etag or ''}|{last_modified or ''}|{filename}|{size}"
                )
            else:
                fingerprint_seed = f"{url_norm}|{content_hash}|{filename}|{size}"
            fingerprint = hashlib.sha256(fingerprint_seed.encode("utf-8")).hexdigest()

            ext_path = cache_dir / f"{fingerprint}{extension}"
            if not ext_path.exists():
                tmp_path.replace(ext_path)
            else:
                tmp_path.unlink(missing_ok=True)

            meta = {
                "fingerprint": fingerprint,
                "url_norm": url_norm,
                "url": url,
                "etag": etag,
                "last_modified": last_modified,
                "content_hash": content_hash,
                "mime_type": content_type.split(";", 1)[0].strip().lower() if content_type else "",
                "size_bytes": size,
                "fetched_at": _now_utc_iso(),
                "text_path": str(ext_path),
            }
            return meta, ext_path, None, False
        except Exception as exc:
            if logger:
                logger.warning(
                    "Attachment download failed (attempt %s/%s) for %s: %s",
                    attempt + 1,
                    config.download_retries + 1,
                    url,
                    exc,
                )
            if attempt >= config.download_retries:
                break
            _retry_delay(attempt)

    return None, None, "download_failed", False


def _extract_text(
    path: Path,
    *,
    mime_type: str,
    config: AttachmentsConfig,
    logger: Optional[logging.Logger] = None,
) -> tuple[str, dict]:
    page_count = 0
    extra_meta: dict | None = None
    if path.suffix.lower() == ".pdf":
        text, page_count = _extract_pdf_text(path)
    elif path.suffix.lower() == ".docx":
        text = _extract_docx_text(path, config=config, logger=logger)
    elif path.suffix.lower() in {".xlsx", ".xls"}:
        text, extra_meta = _extract_spreadsheet_text(path, config=config, logger=logger)
    elif path.suffix.lower() == ".txt":
        text = _extract_txt_text(path)
    elif path.suffix.lower() == ".zip":
        text, extra_meta = _extract_zip_text(path, config=config, logger=logger)
        page_count = int((extra_meta or {}).get("page_count") or 0)
    else:
        return "", {"page_count": 0}

    text = _normalize_extracted_text(text)
    meta = {"page_count": page_count}
    if path.suffix.lower() == ".zip" and extra_meta:
        meta.update({k: v for k, v in extra_meta.items() if k != "page_count"})
    if path.suffix.lower() in {".xlsx", ".xls"} and extra_meta:
        meta.update(extra_meta)

    if config.ocr_enabled and len(text) < config.min_text_chars_for_ocr and path.suffix.lower() == ".pdf":
        ocr_text = _ocr_pdf_text(path, logger=logger)
        if ocr_text:
            text = _normalize_extracted_text(ocr_text)
            meta["ocr_used"] = True
        else:
            meta["ocr_used"] = False
    return text, meta


def _process_attachment_job(
    tender: dict,
    url: str,
    *,
    config: AttachmentsConfig,
    attachment_cache_dir: Path,
    embed_cache_dir: Path,
    embeddings_config: Optional[EmbeddingsConfig],
    retrieval_state: Optional[dict],
    db_path: Path,
    stats: dict,
    llm_state: dict,
    llm_semaphore: Optional[threading.Semaphore],
    tender_state: dict[int, dict],
    stats_lock: threading.Lock,
    tender_lock: threading.Lock,
    logger: Optional[logging.Logger] = None,
) -> None:
    local_store = TenderStore(db_path, logger=logger).open()
    try:
        meta, path, skip_reason, from_cache = _download_attachment(
            url,
            config=config,
            store=local_store,
            cache_dir=attachment_cache_dir,
            logger=logger,
        )
        if skip_reason:
            with stats_lock:
                stats["attachments_skipped"] += 1
            if logger:
                logger.info("Attachment skipped (%s): %s", skip_reason, url)
            with tender_lock:
                metrics = tender_state.setdefault(
                    id(tender),
                    {"summary_fps": set(), "req_keys": set(), "req_norms": set(), "metrics": {}},
                ).setdefault("metrics", {})
                metrics["skipped"] = int(metrics.get("skipped", 0)) + 1
                if skip_reason == "unsupported_type":
                    metrics["unsupported"] = int(metrics.get("unsupported", 0)) + 1
                elif skip_reason == "too_large":
                    metrics["too_large"] = int(metrics.get("too_large", 0)) + 1
                elif skip_reason == "download_failed":
                    metrics["download_failed"] = int(metrics.get("download_failed", 0)) + 1
            return

        if not meta or not path:
            with stats_lock:
                stats["attachments_skipped"] += 1
            with tender_lock:
                metrics = tender_state.setdefault(
                    id(tender),
                    {"summary_fps": set(), "req_keys": set(), "req_norms": set(), "metrics": {}},
                ).setdefault("metrics", {})
                metrics["skipped"] = int(metrics.get("skipped", 0)) + 1
                metrics["download_failed"] = int(metrics.get("download_failed", 0)) + 1
            return

        with stats_lock:
            if from_cache:
                stats["cached"] += 1
            else:
                stats["downloaded"] += 1
            stats["bytes"] += int(meta.get("size_bytes") or 0)

        local_store.upsert_attachment_cache(entry=meta)

        existing = None
        if meta.get("etag") or meta.get("last_modified"):
            existing = local_store.fetch_attachment_cache_by_meta(
                url_norm=meta["url_norm"],
                etag=meta.get("etag"),
                last_modified=meta.get("last_modified"),
            )
        if not existing and meta.get("content_hash"):
            existing = local_store.fetch_attachment_cache_by_hash(
                url_norm=meta["url_norm"],
                content_hash=meta.get("content_hash"),
            )
        force_refresh = "amendment" in (tender.get("_events") or [])
        if force_refresh:
            existing = None

        summary = None
        sections: list[dict] = []
        excerpt_text = ""
        selected_chunks: list[dict] = []
        structured_summary = None
        llm_gated = False
        language = str(getattr(config, "language_default", "en") or "en")
        if existing and existing.get("summary_json"):
            try:
                summary = json.loads(existing["summary_json"])
                if isinstance(summary, dict):
                    structured_summary = summary.get("structured_summary")
                if existing.get("structured_summary_json"):
                    try:
                        structured_summary = json.loads(existing["structured_summary_json"])
                    except Exception:
                        pass
                with stats_lock:
                    stats["llm_cached"] += 1
            except Exception:
                summary = None

        if summary is None:
            text, extract_meta = _extract_text(
                path,
                mime_type=str(meta.get("mime_type") or ""),
                config=config,
                logger=logger,
            )
            if not text:
                with stats_lock:
                    stats["empty_text"] += 1
                    stats["attachments_skipped"] += 1
                if logger:
                    logger.info("Attachment extracted empty text: %s", url)
                with tender_lock:
                    metrics = tender_state.setdefault(
                        id(tender),
                        {"summary_fps": set(), "req_keys": set(), "req_norms": set(), "metrics": {}},
                    ).setdefault("metrics", {})
                    metrics["skipped"] = int(metrics.get("skipped", 0)) + 1
                    metrics["empty_text"] = int(metrics.get("empty_text", 0)) + 1
                    metrics["extraction_failed"] = int(metrics.get("extraction_failed", 0)) + 1
                return

            if config.language_detection_enabled:
                detected = _detect_language(text)
                if detected:
                    language = detected

            if extract_meta.get("ocr_used"):
                with stats_lock:
                    stats["ocr_used"] += 1
            with stats_lock:
                stats["chars"] += len(text)

            zip_files = extract_meta.get("zip_files")
            zip_bytes = extract_meta.get("zip_bytes")
            if zip_files:
                with stats_lock:
                    stats["zip_attachments"] += 1
                    stats["zip_files"] += int(zip_files)
                    if zip_bytes:
                        stats["zip_bytes"] += int(zip_bytes)
                if logger:
                    logger.info(
                        "Attachment zip: files=%s bytes=%s url=%s",
                        zip_files,
                        zip_bytes or 0,
                        url,
                    )

            excerpt, sections = _sectionize(
                text,
                max_chars=config.section_char_budget,
                max_sections=config.max_sections,
            )
            excerpt_text = excerpt
            selected_chunks: list[dict] = []
            if config.chunked_retrieval_enabled and retrieval_state is not None:
                chunk_max_chars = int(getattr(config, "chunk_text_max_chars", 0) or 0)
                chunk_overlap = int(getattr(config, "chunk_text_overlap", 0) or 0)
                chunks = _build_chunks(
                    text,
                    sections,
                    max_chars=chunk_max_chars or 1200,
                    overlap_chars=chunk_overlap or 100,
                )
                max_chunks = int(getattr(config, "chunk_max_per_attachment", 0) or 0)
                if max_chunks <= 0:
                    max_chunks = len(chunks)
                with tender_lock:
                    tender_info = tender_state.setdefault(
                        id(tender),
                        {"summary_fps": set(), "req_keys": set(), "req_norms": set(), "chunk_count": 0},
                    )
                    tender_chunk_count = int(tender_info.get("chunk_count", 0))
                    tender_max = int(getattr(config, "chunk_max_per_tender", 0) or 0)
                    if tender_max > 0:
                        remaining = max(0, tender_max - tender_chunk_count)
                        max_chunks = min(max_chunks, remaining)

                with stats_lock:
                    run_remaining = retrieval_state.get("remaining")
                    if run_remaining is not None:
                        max_chunks = min(max_chunks, int(run_remaining))

                if chunks and max_chunks > 0:
                    scored = _score_chunks(
                        chunks,
                        prompt_embeddings=retrieval_state.get("prompt_embeddings"),
                        embeddings_config=embeddings_config,
                        cache_dir=embed_cache_dir,
                        attachment_key=str(meta.get("fingerprint") or meta.get("url_norm") or url),
                        logger=logger,
                    )
                    selected_chunks = _select_retrieval_chunks(scored, max_chunks)

                if selected_chunks:
                    with stats_lock:
                        run_remaining = retrieval_state.get("remaining")
                        if run_remaining is not None:
                            retrieval_state["remaining"] = max(
                                0, int(run_remaining) - len(selected_chunks)
                            )
                    with tender_lock:
                        tender_info = tender_state.setdefault(
                            id(tender),
                            {"summary_fps": set(), "req_keys": set(), "req_norms": set(), "chunk_count": 0},
                        )
                        tender_info["chunk_count"] = int(tender_info.get("chunk_count", 0)) + len(
                            selected_chunks
                        )
                    chunk_budget = int(getattr(config, "chunk_char_budget", 0) or 0)
                    if chunk_budget <= 0:
                        chunk_budget = config.section_char_budget
                    retrieval_excerpt = _build_retrieval_excerpt(selected_chunks, chunk_budget)
                    if retrieval_excerpt:
                        excerpt = retrieval_excerpt
            if logger:
                logger.info(
                    "Attachment extract: kept=%s chars sections=%s url=%s",
                    len(excerpt),
                    len(sections),
                    url,
                )
            if config.pii_scrub:
                excerpt = _scrub_pii(excerpt)

            llm_gated = False
            if not config.llm_enabled:
                summary = {
                    "requirements": [],
                    "key_risks": [],
                    "dates": [],
                    "attachments_used": [],
                }
                with stats_lock:
                    stats["llm_skipped"] += 1
            else:
                allowed = False
                with stats_lock:
                    remaining = llm_state.get("remaining")
                    daily_cap = llm_state.get("daily_cap", 0)
                    daily_used = llm_state.get("daily_used", 0)
                    if remaining is not None and remaining <= 0:
                        allowed = False
                    elif daily_cap > 0 and daily_used >= daily_cap:
                        allowed = False
                    else:
                        allowed = True
                        if remaining is not None:
                            llm_state["remaining"] = remaining - 1
                        llm_state["daily_used"] = daily_used + 1

                if allowed:
                    try:
                        prompt = _build_prompt(
                            _attachment_name(url), excerpt, language=language
                        )
                        if logger:
                            logger.info("Attachment LLM: requesting summary for %s", url)
                        if llm_semaphore:
                            llm_semaphore.acquire()
                        try:
                            raw_summary = _cohere_chat(prompt, config, logger=logger)
                        finally:
                            if llm_semaphore:
                                llm_semaphore.release()
                        summary = _normalize_summary(raw_summary, evidence_text=excerpt)
                        with stats_lock:
                            stats["llm_requested"] += 1
                            if raw_summary.get("_llm_failed"):
                                stats["llm_skipped"] += 1
                    except Exception:
                        if logger:
                            logger.exception("Attachment LLM failed for %s; skipping summary", url)
                        summary = {
                            "requirements": [],
                            "key_risks": [],
                            "dates": [],
                            "attachments_used": [],
                            "_llm_failed": True,
                        }
                        with stats_lock:
                            stats["llm_skipped"] += 1
                else:
                    summary = {
                        "requirements": [],
                        "key_risks": [],
                        "dates": [],
                        "attachments_used": [],
                        "_llm_failed": True,
                    }
                    with stats_lock:
                        stats["llm_skipped"] += 1
                    llm_gated = True

            if (
                config.structured_summaries_enabled
                and config.llm_enabled
                and structured_summary is None
            ):
                structured_allowed = False
                with stats_lock:
                    remaining = llm_state.get("remaining")
                    daily_cap = llm_state.get("daily_cap", 0)
                    daily_used = llm_state.get("daily_used", 0)
                    if remaining is not None and remaining <= 0:
                        structured_allowed = False
                    elif daily_cap > 0 and daily_used >= daily_cap:
                        structured_allowed = False
                    else:
                        structured_allowed = True
                        if remaining is not None:
                            llm_state["remaining"] = remaining - 1
                        llm_state["daily_used"] = daily_used + 1
                        stats["llm_requested"] += 1

                if structured_allowed:
                    try:
                        attachment_id = str(meta.get("fingerprint") or "")
                        prompt = _build_structured_prompt(
                            _attachment_name(url),
                            attachment_id,
                            excerpt,
                            language=language,
                        )
                        if logger:
                            logger.info(
                                "Attachment LLM: requesting structured summary for %s", url
                            )
                        if llm_semaphore:
                            llm_semaphore.acquire()
                        try:
                            raw_structured = _cohere_chat_structured(
                                prompt, config, logger=logger
                            )
                        finally:
                            if llm_semaphore:
                                llm_semaphore.release()
                        structured_summary = _normalize_structured_summary(
                            raw_structured, attachment_id=attachment_id, evidence_text=excerpt
                        )
                        if isinstance(summary, dict) and structured_summary:
                            summary["structured_summary"] = structured_summary
                    except Exception:
                        if logger:
                            logger.exception(
                                "Attachment structured summary failed for %s", url
                            )

            if config.pii_scrub:
                summary = _scrub_summary_pii(summary)
                if structured_summary:
                    structured_summary = _scrub_structured_summary(structured_summary)

            meta["text_chars"] = len(text)
            meta["summary_cached_at"] = _now_utc_iso()
            if summary.get("_llm_failed"):
                meta["llm_failed"] = True
            meta["summary_json"] = json.dumps(summary, ensure_ascii=False)
            meta["summary_model"] = config.llm_model
            meta["requirements_json"] = json.dumps(summary.get("requirements", []), ensure_ascii=False)
            if summary.get("_llm_failed"):
                meta["summary_failed"] = True
            if structured_summary:
                meta["structured_summary_json"] = json.dumps(
                    structured_summary, ensure_ascii=False
                )
            source_sections = [s.get("heading") for s in sections]
            if source_sections:
                meta["source_sections_json"] = json.dumps(source_sections, ensure_ascii=False)
            if selected_chunks:
                meta["selected_chunks_json"] = json.dumps(
                    [
                        {
                            "heading": chunk.get("heading"),
                            "section_index": chunk.get("section_index"),
                            "chunk_index": chunk.get("chunk_index"),
                            "start_offset": chunk.get("start_offset"),
                            "end_offset": chunk.get("end_offset"),
                            "score": chunk.get("score"),
                            "score_prompt": chunk.get("score_prompt"),
                        }
                        for chunk in selected_chunks
                    ],
                    ensure_ascii=False,
                )
            local_store.upsert_attachment_cache(entry=meta)

        raw_summary = summary or {}
        if isinstance(raw_summary, dict) and structured_summary is None:
            structured_summary = raw_summary.get("structured_summary")
        summary = _normalize_summary(raw_summary, evidence_text=excerpt_text)
        if config.pii_scrub:
            summary = _scrub_summary_pii(summary)
            if structured_summary:
                structured_summary = _scrub_structured_summary(structured_summary)
        source_sections = [s.get("heading") for s in sections]
        memo = _build_attachment_memo(
            summary,
            attachment_name=_attachment_name(url),
            source_sections=source_sections,
        )
        if summary.get("_llm_failed"):
            memo = _excerpt_fallback_memo(
                excerpt_text, attachment_name=_attachment_name(url)
            ) or memo
        attachment_summary = {
            "attachment_url": url,
            "attachment_name": _attachment_name(url),
            "fingerprint": meta.get("fingerprint"),
            "summary": summary,
            "structured_summary": structured_summary,
            "memo": memo,
            "source_sections": source_sections,
            "selected_chunks": [
                {
                    "heading": chunk.get("heading"),
                    "section_index": chunk.get("section_index"),
                    "chunk_index": chunk.get("chunk_index"),
                    "start_offset": chunk.get("start_offset"),
                    "end_offset": chunk.get("end_offset"),
                    "score": chunk.get("score"),
                    "score_prompt": chunk.get("score_prompt"),
                }
                for chunk in selected_chunks
            ],
            "text_chars": meta.get("text_chars"),
            "size_bytes": meta.get("size_bytes"),
            "mime_type": meta.get("mime_type"),
            "language": language,
        }

        requirements = []
        for req in summary.get("requirements", []):
            requirements.append(
                {
                    "requirement": req,
                    "attachment_url": url,
                    "attachment_name": _attachment_name(url),
                    "fingerprint": meta.get("fingerprint"),
                    "source_sections": attachment_summary.get("source_sections", []),
                }
            )

        with tender_lock:
            state = tender_state.setdefault(
                id(tender),
                {"summary_fps": set(), "req_keys": set(), "req_norms": set(), "metrics": {}},
            )
            summary_fps: set[str] = state["summary_fps"]
            req_keys: set[tuple[str, str]] = state["req_keys"]
            req_norms: set[str] = state["req_norms"]
            metrics = state.setdefault("metrics", {})

            fingerprint = str(meta.get("fingerprint") or "")
            if fingerprint and fingerprint not in summary_fps:
                tender.setdefault("attachment_summaries", []).append(attachment_summary)
                summary_fps.add(fingerprint)
            elif not fingerprint:
                tender.setdefault("attachment_summaries", []).append(attachment_summary)

            for req in requirements:
                norm = re.sub(
                    r"\s+", " ", str(req.get("requirement") or "").strip().lower()
                )
                if norm in req_norms:
                    continue
                key = (fingerprint, norm)
                if key in req_keys:
                    continue
                if norm:
                    req_norms.add(norm)
                req_keys.add(key)
                tender.setdefault("attachment_requirements", []).append(req)

            metrics["processed"] = int(metrics.get("processed", 0)) + 1
            metrics["summary_count"] = int(metrics.get("summary_count", 0)) + 1
            if summary.get("requirements"):
                metrics["requirements_extracted"] = int(
                    metrics.get("requirements_extracted", 0)
                ) + 1
            if llm_gated:
                metrics["llm_gated"] = int(metrics.get("llm_gated", 0)) + 1
            if summary.get("_llm_failed"):
                metrics["llm_failed"] = int(metrics.get("llm_failed", 0)) + 1
            raw_reqs = _coerce_list(raw_summary.get("requirements"), max_items=8)
            raw_dates = _coerce_list(raw_summary.get("dates"), max_items=6)
            if (raw_reqs or raw_dates) and not (
                summary.get("requirements") or summary.get("dates")
            ):
                metrics["low_signal"] = int(metrics.get("low_signal", 0)) + 1

        with stats_lock:
            stats["attachments_processed"] += 1
    finally:
        local_store.close()


def enrich_tenders_with_attachments(
    tenders: Iterable[dict],
    *,
    config: AttachmentsConfig,
    store: TenderStore,
    cache_dir: Path,
    embeddings_config: Optional[EmbeddingsConfig] = None,
    logger: Optional[logging.Logger] = None,
) -> dict:
    stats = {
        "attachments_total": 0,
        "attachments_processed": 0,
        "attachments_skipped": 0,
        "downloaded": 0,
        "cached": 0,
        "llm_requested": 0,
        "llm_cached": 0,
        "llm_skipped": 0,
        "ocr_used": 0,
        "empty_text": 0,
        "bytes": 0,
        "chars": 0,
        "zip_attachments": 0,
        "zip_files": 0,
        "zip_bytes": 0,
    }

    if not config.enabled:
        return stats

    if not config.llm_api_key and config.llm_enabled:
        if logger:
            logger.warning("Attachment LLM API key missing; skipping attachment summaries")
        config.llm_enabled = False

    attachment_cache_dir = cache_dir / "attachments"
    attachment_cache_dir.mkdir(parents=True, exist_ok=True)

    today = _today_utc()
    daily_used = 0
    if config.llm_daily_cap > 0:
        usage = _load_daily_usage(attachment_cache_dir, today)
        daily_used = int(usage.get("count", 0))

    remaining_llm = config.llm_per_run_max if config.llm_per_run_max > 0 else None
    llm_state = {
        "remaining": remaining_llm,
        "daily_cap": int(config.llm_daily_cap or 0),
        "daily_used": daily_used,
    }

    retrieval_state: dict | None = None
    if config.chunked_retrieval_enabled:
        prompt_embeddings: dict[str, list[float]] = {}
        if embeddings_config is not None:
            prompt_embeddings, embedded = embed_texts(
                _RETRIEVAL_PROMPTS, embeddings_config, cache_dir, logger=logger
            )
            if not embedded:
                prompt_embeddings = {}
        retrieval_state = {
            "remaining": (
                int(config.chunk_max_per_run)
                if int(getattr(config, "chunk_max_per_run", 0) or 0) > 0
                else None
            ),
            "prompt_embeddings": prompt_embeddings,
        }

    processed_urls = set()
    jobs: list[tuple[dict, str]] = []
    tender_state: dict[int, dict] = {}

    for t in tenders:
        existing_summaries = t.get("attachment_summaries") or []
        existing_summary_fps = {
            s.get("fingerprint") for s in existing_summaries if isinstance(s, dict)
        }
        existing_req_keys = set()
        existing_req_norms = set()
        for req in t.get("attachment_requirements") or []:
            if isinstance(req, dict):
                norm = re.sub(
                    r"\s+", " ", str(req.get("requirement") or "").strip().lower()
                )
                if norm:
                    existing_req_norms.add(norm)
                key = (str(req.get("fingerprint") or ""), norm)
                existing_req_keys.add(key)
        tender_state[id(t)] = {
            "summary_fps": set(existing_summary_fps),
            "req_keys": set(existing_req_keys),
            "req_norms": set(existing_req_norms),
            "metrics": {
                "total": 0,
                "processed": 0,
                "skipped": 0,
                "unsupported": 0,
                "too_large": 0,
                "download_failed": 0,
                "empty_text": 0,
                "extraction_failed": 0,
                "llm_gated": 0,
                "llm_failed": 0,
                "low_signal": 0,
                "summary_count": 0,
                "requirements_extracted": 0,
            },
        }

        attachments = t.get("attachments") or []
        if isinstance(attachments, str):
            attachments = [attachments]
        urls = [a for a in attachments if isinstance(a, str) and a.startswith("http")]
        if not urls:
            continue
        if config.max_attachments_per_tender > 0:
            urls = urls[: config.max_attachments_per_tender]
        tender_state[id(t)]["metrics"]["total"] = len(urls)

        for url in urls:
            if config.max_attachments_per_run > 0 and len(jobs) >= config.max_attachments_per_run:
                break
            if url in processed_urls:
                continue
            processed_urls.add(url)
            jobs.append((t, url))

    stats["attachments_total"] = len(jobs)
    if not jobs:
        return stats

    stats_lock = threading.Lock()
    tender_lock = threading.Lock()
    llm_semaphore = None
    llm_limit = int(getattr(config, "llm_concurrency", 0) or 0)
    if config.llm_enabled and llm_limit > 0:
        llm_semaphore = threading.Semaphore(max(1, llm_limit))
    max_workers = max(1, int(config.download_concurrency or 1))
    db_path = store.db_path

    if max_workers <= 1:
        for t, url in jobs:
            _process_attachment_job(
                t,
                url,
                config=config,
                attachment_cache_dir=attachment_cache_dir,
                embed_cache_dir=cache_dir,
                embeddings_config=embeddings_config,
                retrieval_state=retrieval_state,
                db_path=db_path,
                stats=stats,
                llm_state=llm_state,
                llm_semaphore=llm_semaphore,
                tender_state=tender_state,
                stats_lock=stats_lock,
                tender_lock=tender_lock,
                logger=logger,
            )
    else:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [
                executor.submit(
                    _process_attachment_job,
                    t,
                    url,
                    config=config,
                    attachment_cache_dir=attachment_cache_dir,
                    embed_cache_dir=cache_dir,
                    embeddings_config=embeddings_config,
                    retrieval_state=retrieval_state,
                    db_path=db_path,
                    stats=stats,
                    llm_state=llm_state,
                    llm_semaphore=llm_semaphore,
                    tender_state=tender_state,
                    stats_lock=stats_lock,
                    tender_lock=tender_lock,
                    logger=logger,
                )
                for t, url in jobs
            ]
            for future in as_completed(futures):
                future.result()

    if config.llm_daily_cap > 0:
        _save_daily_usage(attachment_cache_dir, today, llm_state.get("daily_used", daily_used))

    for t in tenders:
        state = tender_state.get(id(t))
        if not state:
            continue
        metrics = dict(state.get("metrics") or {})
        if metrics:
            metrics["summary_exists"] = int(metrics.get("summary_count", 0))
            metrics["requirements_extracted"] = int(metrics.get("requirements_extracted", 0))
            t["attachment_metrics"] = metrics

    return stats


def merge_attachment_requirements(tender: dict) -> None:
    attachment_reqs = tender.get("attachment_requirements") or []
    if not attachment_reqs:
        return

    seen = set()
    merged: list[str] = []
    for req in attachment_reqs:
        text = req.get("requirement") if isinstance(req, dict) else str(req)
        if not text:
            continue
        key = re.sub(r"\s+", " ", text.strip().lower())
        if key in seen:
            continue
        seen.add(key)
        merged.append(text.strip())

    if not merged:
        return

    llm = tender.get("llm")
    if isinstance(llm, dict):
        if "key_requirements_notice" not in llm:
            llm["key_requirements_notice"] = llm.get("key_requirements", [])
        llm["key_requirements"] = merged[:6]
        llm["attachment_summary_used"] = True

    tender["attachment_summary_used"] = True
